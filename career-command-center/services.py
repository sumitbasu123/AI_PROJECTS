from __future__ import annotations

import ipaddress
import hashlib
import os
import re
import socket
import subprocess
import json
from datetime import datetime, timezone, timedelta
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document


DEFAULT_PROFILE = {
    "headline": "Data Engineering Leader",
    "summary": (
        "Add a concise professional summary describing your experience, "
        "leadership scope, and the outcomes you have delivered."
    ),
    "years": 0,
    "team": 0,
    "linkedin_url": "",
    "linkedin_text": "",
    "resume_text": "",
    "roles": [
        "Data Engineering Manager", "Data Platform Architect",
    ],
    "target_companies": [
        "Example Company",
    ],
    "locations": ["India", "Remote", "Kolkata", "Bengaluru", "Hyderabad", "Pune"],
    "skills": [
        "Azure", "Azure Data Factory", "ADF", "Snowflake", "dbt", "Snowpark",
        "Python", "Power BI", "AWS", "IICS", "Airflow", "Spark", "Scala",
        "Hive", "Terraform", "SQL", "Data Architecture", "Data Modeling",
        "Governance", "CI/CD",
    ],
    "proof": [],
}

SKILL_CATALOG = [
    "Azure", "Azure Data Factory", "ADF", "Snowflake", "dbt", "Snowpark",
    "Python", "Power BI", "AWS", "Databricks", "Kafka", "Kubernetes",
    "Terraform", "Spark", "SQL", "Data Modeling", "Governance", "CI/CD",
    "Airflow", "Machine Learning", "GenAI",
]


def normalized(value: object) -> str:
    return str(value or "").lower()


def parse_date(value: str) -> datetime | None:
    if not value:
        return None
    candidate = str(value).strip().replace("Z", "+00:00")
    try:
        parsed = datetime.fromisoformat(candidate)
        return parsed if parsed.tzinfo else parsed.replace(tzinfo=timezone.utc)
    except ValueError:
        return None


def infer_work_mode(text: str, location: str = "") -> str:
    value = normalized(f"{text} {location}")
    if "hybrid" in value:
        return "Hybrid"
    if re.search(r"\bremote\b|work from home|wfh", value):
        return "Remote"
    if re.search(r"\bonsite\b|on-site|office based|in-office", value):
        return "Onsite"
    return "Not specified"


def infer_experience(text: str) -> tuple[int | None, int | None]:
    value = normalized(text)
    ranges = re.findall(
        r"(\d{1,2})\s*(?:\+|to|-|–)\s*(\d{1,2})?\s*(?:years|yrs)",
        value,
    )
    if ranges:
        low, high = ranges[0]
        return int(low), int(high) if high else None
    matches = re.findall(r"(\d{1,2})\+?\s*(?:years|yrs)", value)
    if matches:
        return int(matches[0]), None
    return None, None


def infer_salary(text: str) -> str:
    patterns = [
        r"(?:₹|inr|rs\.?)\s*[\d,.]+\s*(?:-|to|–)\s*(?:₹|inr|rs\.?)?\s*[\d,.]+\s*(?:lpa|lakhs?|pa)?",
        r"\b\d{1,3}\s*(?:-|to|–)\s*\d{1,3}\s*(?:lpa|lakhs?)\b",
        r"(?:usd|\$)\s*[\d,.]+\s*(?:-|to|–)\s*(?:usd|\$)?\s*[\d,.]+",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.I)
        if match:
            return match.group(0)
    return ""


def infer_domain(text: str) -> str:
    value = normalized(text)
    domains = {
        "Healthcare / Life Sciences": [
            "healthcare", "pharma", "life sciences", "patient", "clinical",
            "medical", "gxp", "fda",
        ],
        "Retail / Supply Chain": [
            "retail", "supply chain", "logistics", "manufacturing", "demand planning",
        ],
        "Banking / Financial Services": [
            "banking", "financial services", "payments", "fintech", "insurance",
        ],
        "Technology / SaaS": ["saas", "software platform", "technology company"],
    }
    for domain, terms in domains.items():
        if any(term in value for term in terms):
            return domain
    return "Cross-industry"


def normalize_job(record: dict) -> dict:
    job = dict(record)
    description = clean_html(job.get("description", ""))
    combined = f"{job.get('title', '')} {description}"
    skills = [
        skill for skill in SKILL_CATALOG if normalized(skill) in normalized(combined)
    ]
    cloud_stack = [
        skill for skill in ["Azure", "AWS", "Databricks", "Snowflake"]
        if normalized(skill) in normalized(combined)
    ]
    data_stack = [
        skill for skill in [
            "Azure Data Factory", "ADF", "dbt", "Snowpark", "Python", "Spark",
            "SQL", "Airflow", "Terraform", "Power BI", "Kafka",
        ] if normalized(skill) in normalized(combined)
    ]
    experience_min, experience_max = infer_experience(description)
    now = datetime.now(timezone.utc).isoformat()
    job.update(
        {
            "description": description,
            "work_mode": infer_work_mode(description, job.get("location", "")),
            "experience_min": experience_min,
            "experience_max": experience_max,
            "skills": list(dict.fromkeys(skills)),
            "cloud_stack": list(dict.fromkeys(cloud_stack)),
            "data_stack": list(dict.fromkeys(data_stack)),
            "domain": infer_domain(description),
            "salary": job.get("salary") or infer_salary(description),
            "content_hash": hashlib.sha256(
                re.sub(r"\W+", "", normalized(description)).encode("utf-8")
            ).hexdigest() if description else "",
            "fetched_at": job.get("fetched_at", now),
            "first_seen": job.get("first_seen", now),
            "last_seen": now,
            "raw_payload": job.get("raw_payload", record),
        }
    )
    return job


def credibility_check(job: dict) -> dict:
    score = 70
    flags = []
    source = normalized(job.get("source"))
    text = normalized(f"{job.get('title')} {job.get('description')}")
    url = job.get("url", "")
    parsed = urlparse(url) if url else None

    if re.search(r"company career page|greenhouse|lever|workday|ashby|employer ats", source):
        score += 20
    elif "linkedin alert" in source:
        score += 8
    elif re.search(r"recruiter|consultancy|aggregator", source):
        score -= 15
        flags.append("Third-party source: verify on the employer website")
    if not url:
        score -= 25
        flags.append("No official apply link")
    elif parsed and parsed.scheme not in {"http", "https"}:
        score -= 20
        flags.append("Invalid application link")
    if re.search(r"\burgent\b|immediate joiner|limited openings|apply immediately", text):
        score -= 12
        flags.append("Generic urgency language")
    if re.search(r"whatsapp|telegram|registration fee|processing fee|pay to apply", text):
        score -= 40
        flags.append("Possible recruitment scam language")
    if job.get("salary") and re.search(r"\b\d{3,}\s*lpa\b", normalized(job["salary"])):
        score -= 20
        flags.append("Salary appears unrealistic")
    posted = parse_date(job.get("posted_date", ""))
    if posted:
        age = datetime.now(timezone.utc) - posted.astimezone(timezone.utc)
        if age > timedelta(days=45):
            score -= 20
            flags.append(f"Stale posting: {age.days} days old")
        elif age > timedelta(days=21):
            score -= 8
            flags.append(f"Older posting: {age.days} days old")
    if int(job.get("occurrence_count") or 1) >= 4:
        score -= 10
        flags.append("Frequently reposted")
    return {
        "credibility_score": max(0, min(100, score)),
        "credibility_flags": flags,
    }


def score_job(job: dict, profile: dict) -> dict:
    job = normalize_job(job)
    text = normalized(
        f"{job.get('title', '')} {job.get('description', '')} {job.get('location', '')}"
    )
    title = normalized(job.get("title"))
    source = normalized(job.get("source"))
    matched = [skill for skill in profile["skills"] if normalized(skill) in text]
    skill_score = min(35, round(len(set(matched)) / 8 * 35))

    role_score = 5
    if re.search(r"architect|manager|lead|principal|director", title):
        role_score = 19
    if re.search(r"data engineering manager|data platform architect|data architect|senior manager", title):
        role_score = 25
    if re.search(r"junior|entry level", title):
        role_score = 2

    leadership_score = 4
    if re.search(r"lead|manager|mentor|team|stakeholder|delivery|architecture", text):
        leadership_score = 11
    if re.search(r"lead a team|people manager|engineering manager|director", text):
        leadership_score = 15

    location_score = 6
    if re.search(r"india|remote|kolkata|bengaluru|bangalore|hyderabad|pune|hybrid", text):
        location_score = 15
    if re.search(r"us only|united states only|no sponsorship", text):
        location_score = 2

    credibility = credibility_check(job)
    credibility_score = round(credibility["credibility_score"] / 10)

    risks = list(credibility["credibility_flags"])
    if re.search(r"urgent|immediate joiner", text):
        risks.append("Urgency language")
    if re.search(r"aws glue|kafka|kubernetes|java", text) and not re.search(
        r"azure|snowflake", text
    ):
        risks.append("Primary stack mismatch")
    if "recruiter" in source or "consultancy" in source:
        risks.append("Verify the role on the company website")
    if not job.get("url"):
        risks.append("No official apply link")

    experience_score = 0
    exp_min = job.get("experience_min")
    exp_max = job.get("experience_max")
    if exp_min is None:
        experience_score = 5
    elif exp_min <= profile.get("years", 0) and (
        exp_max is None or profile.get("years", 0) <= exp_max + 5
    ):
        experience_score = 10
    elif exp_min > profile.get("years", 0):
        risks.append(f"Requires {exp_min}+ years")

    score = min(
        100,
        skill_score + role_score + leadership_score + location_score
        + credibility_score + experience_score,
    )
    matched_unique = list(dict.fromkeys(matched))
    reason_parts = []
    if matched_unique:
        reason_parts.append(f"Strong match on {', '.join(matched_unique[:4])}")
    if role_score >= 19:
        reason_parts.append("seniority aligns with architecture/leadership target")
    if location_score == 15:
        reason_parts.append("location or work mode is suitable")
    reason = "; ".join(reason_parts) or "Partial role and profile alignment"
    suggested_action = (
        "Apply" if score >= 82 and credibility["credibility_score"] >= 65
        else "Verify then apply" if score >= 72
        else "Review" if score >= 60
        else "Archive"
    )
    resume_keywords = matched_unique[:6]
    for keyword in ["SAP ingestion", "team leadership", "data governance"]:
        if normalized(keyword) in text and keyword not in resume_keywords:
            resume_keywords.append(keyword)
    return {
        "score": score,
        "matched_skills": matched_unique[:8],
        "risks": risks,
        "keywords": matched_unique[:6],
        "reason": reason,
        "suggested_action": suggested_action,
        "resume_keywords": resume_keywords[:8],
        **credibility,
    }


def analyze_job_description(role: str, company: str, jd: str, profile: dict) -> dict:
    required = [skill for skill in SKILL_CATALOG if normalized(skill) in normalized(jd)]
    matched = [
        skill for skill in required
        if any(
            normalized(skill) == normalized(candidate)
            or normalized(skill) in normalized(candidate)
            for candidate in profile["skills"]
        )
    ]
    gaps = [skill for skill in required if skill not in matched]
    score = score_job(
        {
            "title": role, "company": company, "location": jd,
            "description": jd, "source": "Company career page", "url": "reviewed",
        },
        profile,
    )["score"]
    rounds = [
        "Hiring manager and leadership discussion",
        "Data-platform architecture or system-design round",
        "Technical deep dive on cloud, Snowflake, and data engineering",
        "Behavioral and stakeholder-management round",
    ]
    if re.search(r"coding|python|sql exercise|hands-on", normalized(jd)):
        rounds.insert(2, "SQL or Python practical assessment")
    return {
        "role": role, "company": company, "jd": jd, "required_skills": required,
        "matched_skills": matched, "gaps": gaps, "score": score,
        "likely_rounds": rounds,
    }


def generate_questions(analysis: dict, profile: dict) -> list[dict]:
    role = analysis.get("role") or "target role"
    company = analysis.get("company") or "the company"
    questions = [
        {
            "type": "Behavioral",
            "question": f"Why are you interested in the {role} role at {company}?",
            "guide": "Connect the role to your Azure, Snowflake, delivery, and leadership experience.",
        },
        {
            "type": "Architecture",
            "question": "How would you design a governed SAP S/4HANA to Snowflake data platform on Azure?",
            "guide": "Cover ingestion, layers, orchestration, dbt, security, observability, recovery, and cost.",
        },
        {
            "type": "Technical",
            "question": "When would you choose Snowpark over SQL or dbt models?",
            "guide": "Compare complexity, pushdown, maintainability, deployment, and performance.",
        },
        {
            "type": "Technical",
            "question": "How do you control Snowflake compute cost without harming SLAs?",
            "guide": "Discuss sizing, auto-suspend, isolation, profiling, incremental models, and monitors.",
        },
        {
            "type": "Scenario",
            "question": "A critical pipeline misses its SLA after data volume doubles. What do you do?",
            "guide": "Explain containment, communication, diagnosis, recovery, root cause, and prevention.",
        },
        {
            "type": "Behavioral",
            "question": "Tell me about a technical disagreement with a senior stakeholder.",
            "guide": "Use STAR and show evidence, tradeoffs, ownership, and relationship management.",
        },
        {
            "type": "Architecture",
            "question": "How would you migrate a legacy ETL estate to Azure and Snowflake with limited downtime?",
            "guide": "Cover assessment, waves, dual run, reconciliation, cutover, rollback, and retirement.",
        },
        {
            "type": "Behavioral",
            "question": "How has your leadership style changed while managing 15+ engineers?",
            "guide": "Show delegation, governance, coaching, accountability, and context setting.",
        },
    ]
    skill_guides = {
        "Azure": "Explain service selection, network and identity boundaries, operations, reliability, and cost.",
        "Azure Data Factory": "Cover metadata-driven pipelines, integration runtimes, orchestration, retry, monitoring, and CI/CD.",
        "Snowflake": "Cover workload isolation, security, data organization, performance, cost, sharing, and operations.",
        "dbt": "Discuss model layers, tests, macros, incremental design, documentation, lineage, and deployment.",
        "Snowpark": "Explain why procedural or dataframe logic is justified, execution pushdown, packaging, testing, and cost.",
        "Databricks": "Compare lakehouse patterns, Spark workloads, governance, orchestration, and coexistence with Snowflake.",
        "Terraform": "Cover reusable modules, environment promotion, state, review controls, drift, and rollback.",
        "Governance": "Address ownership, classification, lineage, access controls, quality, retention, and audit.",
        "CI/CD": "Explain branching, automated tests, deployment promotion, secrets, rollback, and release governance.",
    }
    for skill in analysis.get("required_skills", [])[:8]:
        questions.append(
            {
                "type": "Technical",
                "question": (
                    f"How have you used or evaluated {skill} in an enterprise "
                    "data-platform context?"
                ),
                "guide": skill_guides.get(
                    skill,
                    "Explain your real experience, architecture choices, tradeoffs, operations, and outcomes.",
                ),
            }
        )
    questions.extend(
        [
            {
                "type": "Scenario",
                "question": (
                    f"You join {company} and inherit a data platform with rising "
                    "cost, unreliable pipelines, and unclear ownership. What is "
                    "your first 90-day plan?"
                ),
                "guide": (
                    "Structure the answer around discovery, risk control, service "
                    "baselines, architecture priorities, team operating model, "
                    "stakeholder alignment, and measurable milestones."
                ),
            },
            {
                "type": "Behavioral",
                "question": (
                    "Describe how you balance hands-on architecture work with "
                    "project management and leadership of 15+ engineers."
                ),
                "guide": (
                    "Show delegation, architecture guardrails, decision forums, "
                    "coaching, delivery visibility, and when you personally go deep."
                ),
            },
        ]
    )
    for skill in analysis.get("gaps", [])[:4]:
        questions.append(
            {
                "type": "Technical",
                "question": f"What is your practical approach to {skill}?",
                "guide": "Be honest about depth, connect adjacent experience, and give a learning plan.",
            }
        )
    unique_questions = []
    seen = set()
    for question in questions:
        key = normalized(question["question"])
        if key not in seen:
            seen.add(key)
            unique_questions.append(question)
    return unique_questions


def score_answer(answer: str) -> dict:
    words = len(answer.split())
    evidence = bool(re.search(r"\d|percent|team|engineer|project|client|pipeline", answer, re.I))
    ownership = bool(
        re.search(r"\bI (led|designed|architected|decided|owned|implemented|resolved)", answer, re.I)
    )
    tradeoff = bool(re.search(r"tradeoff|because|alternative|risk|decision|however", answer, re.I))
    outcome = bool(re.search(r"result|reduced|improved|saved|delivered|faster|cost|SLA", answer, re.I))
    clarity = min(95, 55 + (20 if words >= 80 else 8) + (10 if "." in answer else 0))
    depth = min(95, 45 + (20 if tradeoff else 0) + (20 if words >= 120 else 8))
    seniority = min(95, 45 + (25 if ownership else 5) + (15 if tradeoff else 0))
    evidence_score = min(95, 35 + (25 if evidence else 0) + (25 if outcome else 0))
    return {
        "overall": round((clarity + depth + seniority + evidence_score) / 4),
        "clarity": clarity, "depth": depth, "seniority": seniority,
        "evidence": evidence_score,
    }


def improve_answer(
    question: str, answer: str, profile: dict, analysis: dict, use_ai: bool = False
) -> str:
    if use_ai and os.getenv("OPENAI_API_KEY"):
        try:
            from openai import OpenAI
            client = OpenAI()
            response = client.responses.create(
                model=os.getenv("OPENAI_MODEL", "gpt-5-mini"),
                input=(
                    "You are a senior interview coach. Do not invent experience. "
                    "Return key feedback, a concise 1-minute answer, and a detailed "
                    "architect-level answer using only supplied facts.\n\n"
                    f"Profile: {profile}\nRole analysis: {analysis}\n"
                    f"Question: {question}\nAnswer: {answer}"
                ),
            )
            return response.output_text
        except Exception as exc:
            return f"AI coaching was unavailable ({exc}).\n\n{rules_based_feedback(answer)}"
    return rules_based_feedback(answer)


def rules_based_feedback(answer: str) -> str:
    improvements = []
    if not re.search(r"\bI (led|designed|architected|decided|owned|implemented)", answer, re.I):
        improvements.append("Make your personal ownership explicit.")
    if not re.search(r"tradeoff|because|alternative|risk|decision", answer, re.I):
        improvements.append("Explain one important tradeoff and why you chose that approach.")
    if not re.search(r"\d|percent|reduced|improved|saved|delivered|SLA", answer, re.I):
        improvements.append("Finish with a measurable or observable outcome.")
    if len(answer.split()) < 80:
        improvements.append("Add implementation detail suitable for an architect-level answer.")
    if not improvements:
        improvements.append("Strong structure. Tighten it into a clear 60-90 second version.")
    return "\n".join(f"- {item}" for item in improvements)


def extract_docx_text(source) -> str:
    document = Document(source)
    blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(blocks)


def clean_html(value: str) -> str:
    soup = BeautifulSoup(value or "", "html.parser")
    return re.sub(r"\s+", " ", soup.get_text(" ", strip=True)).strip()


def get_json(url: str, params: dict | None = None, timeout: int = 25) -> dict:
    if os.name == "nt":
        prepared = requests.Request("GET", url, params=params).prepare().url
        command = [
            "powershell.exe",
            "-NoProfile",
            "-NonInteractive",
            "-Command",
            (
                "$ProgressPreference='SilentlyContinue'; "
                "$OutputEncoding=[Console]::OutputEncoding="
                "[System.Text.UTF8Encoding]::new(); "
                "Invoke-RestMethod -Uri $env:CAREER_API_URL -TimeoutSec 30 | "
                "ConvertTo-Json -Depth 20 -Compress"
            ),
        ]
        environment = os.environ.copy()
        environment["CAREER_API_URL"] = prepared
        result = subprocess.run(
            command, capture_output=True, text=True, encoding="utf-8",
            errors="replace", timeout=timeout + 10, check=True, env=environment,
        )
        return json.loads(result.stdout)

    headers = {
        "Accept": "application/json",
        "User-Agent": "CareerCommandCenter/2.0 personal-job-search",
    }
    try:
        response = requests.get(
            url, params=params, headers=headers, timeout=timeout
        )
        response.raise_for_status()
        return response.json()
    except requests.RequestException:
        raise


def build_search_queries(profile: dict, resume_text: str = "") -> list[str]:
    roles = profile.get("roles", [])[:4]
    preferred = [
        "Azure", "Snowflake", "dbt", "Snowpark", "Data Architecture",
        "Data Engineering",
    ]
    available = [
            skill for skill in preferred
            if skill.lower() in normalized(
            f"{resume_text} {profile.get('linkedin_text', '')} "
            f"{' '.join(profile.get('skills', []))}"
        )
    ]
    queries = list(roles)
    if available:
        queries.append(" ".join(available[:2]) + " Architect")
    primary_role = roles[0] if roles else "Data Engineering Manager"
    for company in profile.get("target_companies", [])[:2]:
        queries.append(f"{company} {primary_role}")
    return list(dict.fromkeys(query.strip() for query in queries if query.strip()))[:7]


def search_adzuna(
    query: str,
    location: str,
    app_id: str,
    app_key: str,
    country: str = "in",
    results_per_page: int = 25,
) -> list[dict]:
    if not app_id or not app_key:
        return []
    payload = get_json(
        f"https://api.adzuna.com/v1/api/jobs/{country}/search/1",
        params={
            "app_id": app_id,
            "app_key": app_key,
            "what": query,
            "where": location,
            "results_per_page": results_per_page,
            "content-type": "application/json",
            "sort_by": "date",
        },
        timeout=25,
    )
    jobs = []
    for item in payload.get("results", []):
        jobs.append(
            {
                "external_id": f"adzuna:{item.get('id', '')}",
                "title": item.get("title", "").strip(),
                "company": (item.get("company") or {}).get("display_name", "Unknown"),
                "location": (item.get("location") or {}).get("display_name", location),
                "source": "Adzuna",
                "url": item.get("redirect_url", ""),
                "description": clean_html(item.get("description", "")),
                "posted_date": item.get("created", ""),
            }
        )
    return jobs


def search_remotive(query: str, limit: int = 50) -> list[dict]:
    payload = get_json(
        "https://remotive.com/api/remote-jobs",
        params={"search": query, "limit": limit},
        timeout=25,
    )
    jobs = []
    for item in payload.get("jobs", []):
        jobs.append(
            {
                "external_id": f"remotive:{item.get('id', '')}",
                "title": item.get("title", "").strip(),
                "company": item.get("company_name", "Unknown"),
                "location": item.get("candidate_required_location") or "Remote",
                "source": "Remotive",
                "url": item.get("url", ""),
                "description": clean_html(item.get("description", "")),
                "posted_date": item.get("publication_date", ""),
            }
        )
    return jobs


def search_arbeitnow(query: str, pages: int = 2) -> list[dict]:
    terms = [
        term for term in re.findall(r"[a-zA-Z0-9+#.]+", query.lower())
        if len(term) > 2 and term not in {"and", "the", "with"}
    ]
    jobs = []
    next_url = "https://www.arbeitnow.com/api/job-board-api"
    for _ in range(max(1, min(pages, 5))):
        payload = get_json(next_url, timeout=25)
        for item in payload.get("data", []):
            haystack = normalized(
                f"{item.get('title', '')} {item.get('description', '')} "
                f"{' '.join(item.get('tags') or [])}"
            )
            if terms and not any(term in haystack for term in terms):
                continue
            location = item.get("location", "")
            if item.get("remote"):
                location = f"{location} / Remote" if location else "Remote"
            jobs.append(
                {
                    "external_id": f"arbeitnow:{item.get('slug', '')}",
                    "title": item.get("title", "").strip(),
                    "company": item.get("company_name", "Unknown"),
                    "location": location,
                    "source": "Arbeitnow / employer ATS",
                    "url": item.get("url", ""),
                    "description": clean_html(item.get("description", "")),
                    "posted_date": datetime.fromtimestamp(
                        item.get("created_at", 0), tz=timezone.utc
                    ).isoformat() if item.get("created_at") else "",
                }
            )
        next_url = (payload.get("links") or {}).get("next")
        if not next_url:
            break
    return jobs


def discover_jobs(
    profile: dict,
    resume_text: str,
    location: str,
    providers: list[str],
    adzuna_app_id: str = "",
    adzuna_app_key: str = "",
    minimum_score: int = 65,
    max_queries: int = 5,
) -> dict:
    queries = build_search_queries(profile, resume_text)[:max_queries]
    collected = []
    errors = []
    for query in queries:
        provider_calls = []
        if "Adzuna India" in providers:
            provider_calls.append(
                ("Adzuna India", search_adzuna, (query, location, adzuna_app_id, adzuna_app_key))
            )
        if "Remotive" in providers:
            provider_calls.append(("Remotive", search_remotive, (query,)))
        if "Arbeitnow" in providers:
            provider_calls.append(("Arbeitnow", search_arbeitnow, (query,)))
        for provider, function, arguments in provider_calls:
            try:
                collected.extend(function(*arguments))
            except Exception as exc:
                errors.append(f"{provider}: {exc}")

    deduplicated = {}
    description_companies = {}
    for job in collected:
        job = normalize_job(job)
        identity = normalized(
            f"{job.get('company')}|{job.get('title')}|{job.get('location')}"
        )
        key = job.get("url") or identity
        if not key or not job.get("title") or len(job.get("description", "")) < 80:
            continue
        content_hash = job.get("content_hash")
        if content_hash:
            description_companies.setdefault(content_hash, set()).add(
                normalized(job.get("company"))
            )
        analysis = score_job(job, profile)
        job["analysis"] = analysis
        job["status"] = (
            "Shortlisted"
            if analysis["suggested_action"] == "Apply"
            else "Review"
        )
        job["fetched_at"] = datetime.now(timezone.utc).isoformat()
        if analysis["score"] >= minimum_score:
            previous = deduplicated.get(key)
            if not previous or analysis["score"] > previous["analysis"]["score"]:
                deduplicated[key] = job
    for job in deduplicated.values():
        companies = description_companies.get(job.get("content_hash"), set())
        if len(companies) >= 3:
            job["analysis"]["credibility_score"] = max(
                0, job["analysis"]["credibility_score"] - 20
            )
            job["analysis"]["credibility_flags"].append(
                "Description copied across multiple companies"
            )
            job["analysis"]["risks"].append(
                "Description copied across multiple companies"
            )
            if job["analysis"]["suggested_action"] == "Apply":
                job["analysis"]["suggested_action"] = "Verify then apply"
    jobs = sorted(
        deduplicated.values(),
        key=lambda item: item["analysis"]["score"],
        reverse=True,
    )
    return {"jobs": jobs, "queries": queries, "errors": errors}


def parse_linkedin_alert_text(text: str) -> list[dict]:
    """Parse pasted LinkedIn alert email text without accessing LinkedIn."""
    if not text.strip():
        return []
    url_pattern = re.compile(
        r"https?://(?:(?:[a-z]{2}|www)\.)?linkedin\.com/jobs/view/[^\s<>\"]+",
        re.I,
    )
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    jobs = []
    for index, line in enumerate(lines):
        urls = url_pattern.findall(line)
        if not urls:
            continue
        url = urls[0].rstrip(").,]")
        context = lines[max(0, index - 4):index + 2]
        candidates = [
            item for item in context
            if not url_pattern.search(item)
            and len(item) < 180
            and not re.search(r"view jobs|job alert|unsubscribe|linkedin", item, re.I)
        ]
        title = candidates[0] if candidates else "LinkedIn alert job"
        company = candidates[1] if len(candidates) >= 2 else "Company not parsed"
        location = candidates[2] if len(candidates) >= 3 else ""
        jobs.append(
            normalize_job(
                {
                    "external_id": f"linkedin-alert:{hashlib.sha256(url.encode()).hexdigest()[:16]}",
                    "title": title,
                    "company": company,
                    "location": location,
                    "source": "LinkedIn alert email",
                    "url": url,
                    "description": " ".join(context),
                    "raw_payload": {"alert_text": " ".join(context)},
                }
            )
        )
    unique_jobs = {}
    for job in jobs:
        unique_jobs[job["url"]] = job
    return list(unique_jobs.values())


def resume_recommendations(job: dict, profile: dict) -> dict:
    analysis = score_job(job, profile)
    required = [
        skill for skill in SKILL_CATALOG
        if normalized(skill) in normalized(job.get("description"))
    ]
    gaps = [
        skill for skill in required
        if not any(normalized(skill) in normalized(candidate) for candidate in profile["skills"])
    ]
    proof = profile.get("proof", [])
    bullets = []
    for item in proof:
        relevant = [
            keyword for keyword in analysis["resume_keywords"]
            if normalized(keyword.split()[0]) in normalized(item)
        ]
        if relevant:
            bullets.append(item)
    if not bullets:
        bullets = proof[:3]
    return {
        "headline": f"{job.get('title')} | Azure, Snowflake, dbt & Data Platform Leadership",
        "keywords": analysis["resume_keywords"],
        "proof_points": bullets[:4],
        "gaps": gaps[:6],
        "cover_note": (
            f"I am interested in the {job.get('title')} opportunity at "
            f"{job.get('company')}. My background combines enterprise data "
            "architecture, hands-on Azure/Snowflake delivery, and leadership "
            "of a 15+ member engineering team. I would welcome the opportunity "
            "to discuss how this experience can support the role's platform "
            "and delivery objectives."
        ),
    }


def build_daily_summary(jobs) -> str:
    if hasattr(jobs, "to_dict"):
        rows = jobs.to_dict("records")
    else:
        rows = list(jobs)
    today = datetime.now().strftime("%B %d, %Y")
    active = [row for row in rows if row.get("status") != "Archive"]
    ranked = sorted(active, key=lambda row: int(row.get("fit_score") or 0), reverse=True)
    lines = [
        f"# Career Command Daily Summary - {today}",
        "",
        f"- Active opportunities: {len(active)}",
        f"- Shortlisted: {sum(row.get('status') == 'Shortlisted' for row in active)}",
        f"- Applied/interview: {sum(row.get('status') in {'Applied', 'Interview'} for row in active)}",
        "",
        "## Top Opportunities",
        "",
    ]
    for row in ranked[:10]:
        lines.extend(
            [
                f"### {row.get('title')} - {row.get('company')}",
                f"- Fit: {row.get('fit_score', 0)}/100",
                f"- Status: {row.get('status', 'Review')}",
                f"- Location: {row.get('location') or 'Not specified'}",
                f"- Suggested action: {row.get('suggested_action') or 'Review'}",
                f"- Apply: {row.get('url') or 'No official link'}",
                "",
            ]
        )
    return "\n".join(lines)


def fetch_job_page(url: str) -> dict:
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.hostname:
        raise ValueError("Enter a valid HTTP or HTTPS job URL.")
    try:
        addresses = socket.getaddrinfo(parsed.hostname, parsed.port or 443)
        for address in addresses:
            ip = ipaddress.ip_address(address[4][0])
            if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved:
                raise ValueError("Local or private network URLs are not allowed.")
    except socket.gaierror as exc:
        raise ValueError("The job-site hostname could not be resolved.") from exc

    response = requests.get(
        url,
        headers={"User-Agent": "CareerCommandCenter/1.0 personal-job-research"},
        timeout=15,
    )
    response.raise_for_status()
    if len(response.content) > 3_000_000:
        raise ValueError("The page is too large to import safely.")
    soup = BeautifulSoup(response.text, "html.parser")
    for element in soup(["script", "style", "noscript", "svg"]):
        element.decompose()
    title = soup.title.get_text(" ", strip=True) if soup.title else parsed.hostname
    text = re.sub(r"\s+", " ", soup.get_text(" ", strip=True))
    return {"title": title[:200], "text": text[:30_000], "url": url}
